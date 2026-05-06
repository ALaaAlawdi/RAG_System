from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
import chromadb
from langchain_core.documents import Document
import docx
import httpx
from html.parser import HTMLParser
from ..core.logger import setup_logger
from typing import Optional, List

logger = setup_logger(__name__)


async def get_openai_embedding_model() -> OpenAIEmbeddings:
    return OpenAIEmbeddings(model="text-embedding-3-large")


async def get_persistent_client():
    return chromadb.PersistentClient(path="./chroma_db")


async def get_vector_store(collection_name: Optional[str]) -> Chroma:
    embedding_model = await get_openai_embedding_model()
    client = await get_persistent_client()
    return Chroma(
        client=client,
        collection_name=collection_name if collection_name else "base_center",
        embedding_function=embedding_model,
    )


async def get_text_splitter() -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=100,
        separators=["\n\n", "\n", ".", " "],
    )


async def get_text_chunks(text: str) -> Optional[list[str]]:
    splitter = await get_text_splitter()
    return splitter.split_text(text)



async def read_txt(file_path: str) -> List[Document]:
    documents = []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        chunks = await get_text_chunks(text)
        for i, chunk in enumerate(chunks):
            documents.append(
                Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "page": 1,
                        "chunk": i,
                    },
                )
            )
        return documents
    except Exception as e:
        logger.error(f"Error reading {file_path}: {e}")
        return []


async def read_docx(file_path: str) -> List[Document]:
    try:
        doc = docx.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        chunks = await get_text_chunks(text)
        return [
            Document(page_content=chunk, metadata={"source": file_path, "page": 1, "chunk": i})
            for i, chunk in enumerate(chunks)
        ]
    except Exception as e:
        logger.error(f"Error reading {file_path}: {e}")
        return []


async def read_markdown(file_path: str) -> List[Document]:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        chunks = await get_text_chunks(text)
        return [
            Document(page_content=chunk, metadata={"source": file_path, "page": 1, "chunk": i})
            for i, chunk in enumerate(chunks)
        ]
    except Exception as e:
        logger.error(f"Error reading {file_path}: {e}")
        return []


class _TextExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self._parts: list[str] = []
        self._skip_tags = {"script", "style", "head", "nav", "footer"}
        self._active_skip = 0

    def handle_starttag(self, tag, _attrs):
        if tag in self._skip_tags:
            self._active_skip += 1

    def handle_endtag(self, tag):
        if tag in self._skip_tags:
            self._active_skip = max(0, self._active_skip - 1)

    def handle_data(self, data):
        if self._active_skip == 0 and data.strip():
            self._parts.append(data.strip())

    def get_text(self) -> str:
        return "\n".join(self._parts)


async def fetch_url(url: str) -> List[Document]:
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=30) as client:
            response = await client.get(url)
            response.raise_for_status()
        parser = _TextExtractor()
        parser.feed(response.text)
        text = parser.get_text()
        chunks = await get_text_chunks(text)
        return [
            Document(page_content=chunk, metadata={"source": url, "page": 1, "chunk": i})
            for i, chunk in enumerate(chunks)
        ]
    except Exception as e:
        logger.error(f"Error fetching {url}: {e}")
        return []